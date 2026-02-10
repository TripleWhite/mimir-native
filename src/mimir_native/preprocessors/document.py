"""
Document Preprocessor
支持 PDF、DOCX、TXT 文件的文本提取和结构化处理
"""

import logging
from pathlib import Path
from typing import Any, Optional, List, Dict
from datetime import datetime

from .base import BasePreprocessor, RawContent, parse_date

logger = logging.getLogger(__name__)

# 可选依赖导入
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    logger.warning("PyPDF2 未安装，PDF 处理将不可用")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx 未安装，DOCX 处理将不可用")


class DocumentProcessor(BasePreprocessor):
    """
    文档预处理器
    
    支持的格式：
    - PDF (.pdf)
    - Word (.docx)
    - 文本 (.txt)
    """
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.md', '.rst'}
    
    def supports(self, content_type: str) -> bool:
        """检查是否支持指定的内容类型"""
        return content_type.lower() in {'document', 'pdf', 'docx', 'txt', 'text'}
    
    def process(self, content: Any, metadata: dict) -> RawContent:
        """
        处理文档内容
        
        Args:
            content: 文件路径 (str/Path) 或文件内容 (bytes/str)
            metadata: 元数据字典，可包含 file_path, file_name 等
            
        Returns:
            RawContent: 标准化的内容对象
        """
        try:
            # 确定内容类型
            file_path = metadata.get('file_path', '')
            file_name = metadata.get('file_name', '')
            
            if isinstance(content, (str, Path)):
                file_path = str(content)
                content_type = Path(file_path).suffix.lower()
                text = self._extract_from_file(file_path, content_type)
            elif isinstance(content, bytes):
                content_type = metadata.get('content_type', '.txt')
                text = self._extract_from_bytes(content, content_type, file_name)
            else:
                text = str(content)
                content_type = '.txt'
            
            # 提取文档元数据
            doc_metadata = self._extract_metadata(file_path, text)
            
            # 分块处理
            chunks = self._chunk_by_structure(text)
            
            # 生成摘要
            summary = self._generate_summary(text)
            
            # 解析发生时间
            occurred_at = parse_date(metadata.get('created_date') or metadata.get('modified_date'))
            
            return RawContent(
                text=text,
                summary=summary,
                chunks=chunks,
                metadata={
                    'file_path': file_path,
                    'file_name': file_name or Path(file_path).name if file_path else '',
                    'file_type': content_type,
                    **doc_metadata,
                    **metadata
                },
                occurred_at=occurred_at
            )
            
        except Exception as e:
            logger.error(f"文档处理失败: {e}")
            # 返回错误信息作为内容
            return RawContent(
                text=f"[文档处理错误] {str(e)}",
                summary="文档处理失败",
                chunks=[],
                metadata={'error': str(e), **metadata},
                occurred_at=None
            )
    
    def _extract_from_file(self, file_path: str, content_type: str) -> str:
        """从文件路径提取文本"""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        suffix = path.suffix.lower()
        
        if suffix == '.pdf':
            return self._extract_pdf(file_path)
        elif suffix == '.docx':
            return self._extract_docx(file_path)
        elif suffix in {'.txt', '.md', '.rst', ''}:
            return self._extract_text(file_path)
        else:
            # 尝试作为文本读取
            return self._extract_text(file_path)
    
    def _extract_from_bytes(self, content: bytes, content_type: str, file_name: str = '') -> str:
        """从字节内容提取文本"""
        suffix = Path(file_name).suffix.lower() if file_name else content_type.lower()
        
        if suffix == '.pdf':
            if not PYPDF2_AVAILABLE:
                raise ImportError("PyPDF2 未安装，无法处理 PDF")
            return self._extract_pdf_bytes(content)
        elif suffix == '.docx':
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx 未安装，无法处理 DOCX")
            return self._extract_docx_bytes(content)
        else:
            # 尝试作为文本解码
            try:
                return content.decode('utf-8')
            except UnicodeDecodeError:
                return content.decode('latin-1')
    
    def _extract_pdf(self, file_path: str) -> str:
        """提取 PDF 文本"""
        if not PYPDF2_AVAILABLE:
            raise ImportError("PyPDF2 未安装，无法处理 PDF")
        
        text_parts = []
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                num_pages = len(pdf_reader.pages)
                
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[Page {i+1}/{num_pages}]\n{page_text}")
                    
        except Exception as e:
            logger.error(f"PDF 提取失败: {e}")
            raise
        
        return "\n\n".join(text_parts)
    
    def _extract_pdf_bytes(self, content: bytes) -> str:
        """从字节提取 PDF 文本"""
        import io
        text_parts = []
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            num_pages = len(pdf_reader.pages)
            
            for i, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"[Page {i+1}/{num_pages}]\n{page_text}")
                    
        except Exception as e:
            logger.error(f"PDF 字节提取失败: {e}")
            raise
        
        return "\n\n".join(text_parts)
    
    def _extract_docx(self, file_path: str) -> str:
        """提取 DOCX 文本"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 未安装，无法处理 DOCX")
        
        text_parts = []
        try:
            doc = docx.Document(file_path)
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
                    
        except Exception as e:
            logger.error(f"DOCX 提取失败: {e}")
            raise
        
        return "\n".join(text_parts)
    
    def _extract_docx_bytes(self, content: bytes) -> str:
        """从字节提取 DOCX 文本"""
        import io
        text_parts = []
        try:
            doc_file = io.BytesIO(content)
            doc = docx.Document(doc_file)
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
                    
        except Exception as e:
            logger.error(f"DOCX 字节提取失败: {e}")
            raise
        
        return "\n".join(text_parts)
    
    def _extract_text(self, file_path: str) -> str:
        """提取纯文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
    
    def _extract_metadata(self, file_path: str, text: str) -> Dict[str, Any]:
        """提取文档元数据"""
        metadata = {}
        
        if file_path:
            path = Path(file_path)
            if path.exists():
                stat = path.stat()
                metadata['file_size'] = stat.st_size
                metadata['modified_time'] = datetime.fromtimestamp(stat.st_mtime).isoformat()
        
        # 统计信息
        metadata['char_count'] = len(text)
        metadata['word_count'] = len(text.split())
        metadata['line_count'] = len(text.splitlines())
        
        return metadata
    
    def _chunk_by_structure(self, text: str) -> List[str]:
        """
        按文档结构分块
        
        策略：
        1. 优先按标题分块
        2. 其次按段落分块
        3. 最后按固定大小分块
        """
        chunks = []
        lines = text.split('\n')
        current_chunk = []
        
        for line in lines:
            stripped = line.strip()
            
            # 检测标题行（简单启发式）
            is_heading = (
                stripped and
                (stripped.startswith('#') or  # Markdown
                 stripped.startswith('===') or stripped.startswith('---') or  # 下划线标题
                 (len(stripped) < 100 and stripped.endswith(':')) or  # 冒号结尾的短行
                 (len(current_chunk) == 0 and len(stripped) < 80 and stripped.isupper())  # 全大写标题
                )
            )
            
            if is_heading and current_chunk:
                # 保存当前块并开始新块
                chunk_text = '\n'.join(current_chunk).strip()
                if chunk_text:
                    chunks.append(chunk_text)
                current_chunk = [line]
            else:
                current_chunk.append(line)
        
        # 添加最后一块
        if current_chunk:
            chunk_text = '\n'.join(current_chunk).strip()
            if chunk_text:
                chunks.append(chunk_text)
        
        # 如果分块太大，进一步分割
        final_chunks = []
        for chunk in chunks:
            if len(chunk) > 2000:
                final_chunks.extend(self._chunk_text(chunk, chunk_size=1500, overlap=100))
            else:
                final_chunks.append(chunk)
        
        return final_chunks if final_chunks else self._chunk_text(text)
